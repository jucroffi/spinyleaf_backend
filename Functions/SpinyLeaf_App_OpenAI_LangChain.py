
import os
import sys
import openai
import tiktoken
from openai import OpenAI
from openai import AsyncOpenAI


from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from pathlib import Path

import pandas as pd
import numpy as np
import re

from pathlib import Path
import re

from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from langchain.chat_models import ChatOpenAI




try:
    openai.api_key = sys.argv[1]
except IndexError:
    print("OpenAI API key not provided.")
    sys.exit(1)

key = openai.api_key
client = OpenAI(api_key= key) 



main_f = Path.home() / "SpinyLeaf_App" 
out_f = main_f / "Wellbeing_Fostered_by_Design"

wb_path = out_f / 'Wellbeing.csv'
com_path = out_f / 'Comfort_Dimension' / 'Comfort.csv'
mat_path = out_f / 'Comfort_Dimension' / 'Materials.csv'
del_path = out_f / 'Delight_Dimension' / 'Delight.csv'
soc_path = out_f / 'Social_Dimension' / 'Social.csv'

wb_im_path = out_f 
com_im_path = out_f / 'Comfort_Dimension'
del_im_path = out_f / 'Delight_Dimension'
soc_im_path = out_f / 'Social_Dimension'



def get_completion(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=0.3,
    )
    return response.choices[0].message.content



def create_rag_chunks(text):
    """
    Splits a structured reference text into RAG-friendly chunks.
    Each chunk is a paragraph or bullet group.
    """
    
    raw_chunks = text.split("\n\n")

    chunks = [chunk.strip() for chunk in raw_chunks if chunk.strip()]

    return chunks


def format_llama3_prompt(user_prompt):
    return (
        "<|begin_of_text|>"
        "<|start_header_id|>user<|end_header_id|>\n"
        f"{user_prompt}\n"
        "<|start_header_id|>assistant<|end_header_id|>\n"
    )

def df_to_dict(path, cols=None, filter_col=None, filter_val=None, exclude_prefixes=None):
    df = pd.read_csv(path)

    if cols:
        df = df[cols]

    if filter_col and filter_val is not None:
        df = df[df[filter_col] < filter_val]

    if exclude_prefixes:
        pattern = '|'.join(f"^{prefix}" for prefix in exclude_prefixes)
        df = df[~df["room_ids"].str.contains(pattern, regex=True)]

    return df.to_dict(orient="records")


def extract_wellbeing_summary(path):
    
    df = pd.read_csv(path)

    avg = round(df["wellbeing_satisfaction"].mean(), 2)

    satisfied = df[df["wellbeing_satisfaction"] >= 4].shape[0]
    neutral = df[(df["wellbeing_satisfaction"] >= 2) & (df["wellbeing_satisfaction"] < 4)].shape[0]
    dissatisfied = df[df["wellbeing_satisfaction"] < 2].shape[0]

    low_rooms_df = df[df["wellbeing_satisfaction"] < 3][["room_ids", "wellbeing_satisfaction"]]
    low_rooms = ', '.join(
        f'{row.room_ids} ({round(row.wellbeing_satisfaction, 2)})'
        for _, row in low_rooms_df.iterrows()
    ) or "None"

    mean_comfort = round(df["comfort_satisfaction"].mean(), 2)
    mean_delight = round(df["delight_satisfaction"].mean(), 2)
    mean_social = round(df["social_satisfaction"].mean(), 2)

    return avg, satisfied, neutral, dissatisfied, low_rooms, mean_comfort, mean_delight, mean_social


def extract_comfort_summary(path):
    df = pd.read_csv(path)

    means = {
        "thermal": round(df["extreme_hot_satisf"].mean(), 2),
        "daylight": round(df["daylight_satisf"].mean(), 2),
        "acoustic": round(df["sound_satisf"].mean(), 2),
        "air_quality": round(df["air_quali"].mean(), 2)
    }


    issues = []
    for factor, col in {
        "thermal": ["extreme_hot_satisf"],
        "daylight": ["daylight_satisf"],
        "acoustic": ["sound_satisf"],
        "air_quality": ["air_quali"]
    }.items():
        for c in col:
            df_issue = df[df[c] < 0.66][["room_ids", "floor_level", c]]
            for _, row in df_issue.iterrows():
                issues.append(f"- {row.room_ids} (Floor {row.floor_level}): {factor} = {round(row[c], 2)}")


    factor_issue_counts = {
        "thermal": (df["extreme_hot_satisf"] < 0.66).sum(),
        "daylight": (df["daylight_satisf"] < 0.66).sum(),
        "acoustic": (df["sound_satisf"] < 0.66).sum(),
        "air_quality": (df["air_quali"] < 0.66).sum()
    }


    max_issues = max(factor_issue_counts.values())
    worst_factors = [factor for factor, count in factor_issue_counts.items() if count == max_issues and count > 0]

    return means, issues, factor_issue_counts, worst_factors


def extract_delight_summary(path):
    df = pd.read_csv(path)

    means = {
        "views": round(df["views_overall_satisf"].mean(), 2),
        "balcony": round(df["balcony_satisf"].mean(), 2),
        "space_size": round(df["space_size_satisf"].mean(), 2)
    }

    issues = []
    factor_columns = {
        "views": "views_overall_satisf",
        "balcony": "balcony_satisf",
        "space_size": "space_size_satisf"
    }

    for factor, column in factor_columns.items():
        df_issue = df[(df[column] < 0.66) & (~df["room_ids"].str.startswith("CORE"))][["room_ids", "floor_level", column]]
        for _, row in df_issue.iterrows():
            issues.append(f"- {row.room_ids} (Floor {row.floor_level}): {factor} = {round(row[column], 2)}")

    factor_issue_counts = {
        factor: ((df[column] < 0.66) & (~df["room_ids"].str.startswith("CORE"))).sum()
        for factor, column in factor_columns.items()
    }

    max_issues = max(factor_issue_counts.values())
    worst_factors = [factor for factor, count in factor_issue_counts.items() if count == max_issues and count > 0]

    return means, issues, factor_issue_counts, worst_factors


def extract_social_summary(path):

    df = pd.read_csv(path)

    means = {
        "social_amount": round(df["social_amount_satisf"].mean(), 2),
        "social_distribution": round(df["social_distribution_satisf"].mean(), 2),
        "social_green": round(df["social_green_satisf"].mean(), 2)
    }


    issues = []
    factor_columns = {
        "social_amount": "social_amount_satisf",
        "social_distribution": "social_distribution_satisf",
        "social_green": "social_green_satisf"
    }

    for factor, column in factor_columns.items():
        df_issue = df[df[column] < 0.66][["room_ids", "floor_level", column]]
        for _, row in df_issue.iterrows():
            issues.append(f"- {row.room_ids} (Floor {row.floor_level}): {factor} = {round(row[column], 2)}")

    factor_issue_counts = {
        factor: (df[column] < 0.66).sum()
        for factor, column in factor_columns.items()
    }


    max_issues = max(factor_issue_counts.values())
    worst_factors = [factor for factor, count in factor_issue_counts.items() if count == max_issues and count > 0]

    return means, issues, factor_issue_counts, worst_factors


def extract_materials_summary_enhanced(path):
    df = pd.read_csv(path)
    row = df.iloc[0]

    window_type_str = row["Window_type"]
    glazing_match = re.search(r'(Sgl|Dbl|Trp)', window_type_str)
    glazing = {
        'Sgl': 'Single glazing',
        'Dbl': 'Double glazing',
        'Trp': 'Triple glazing'
    }.get(glazing_match.group(1), 'Unknown glazing') if glazing_match else 'Unknown glazing'

    glass_type = window_type_str.split(glazing_match.group(1))[1].strip() if glazing_match else window_type_str

    summary = {
        "glazing_type": glazing,
        "glass_type": glass_type,
        "windows_u": round(row["Windows_U"], 2),
        "window_noise_reduction": round(row["Win_reduction"], 1),
        "shgc": round(row["SHGC"], 2),
        "wall_r_insulation": row["Wall_R"],
        "wall_noise_reduction": row["Wall_reducrion"],
        "roof_r_insulation": row["Roof_R"],
        "ground_r_insulation": row["Ground_R"]
    }

    return summary





exclude = ["CORE", "SOCIAL"]

wellbeing_df = df_to_dict(
    wb_path,
    cols=["room_ids", "floor_area", "wellbeing_satisfaction", "comfort_satisfaction", "delight_satisfaction", "social_satisfaction"],
    exclude_prefixes=exclude
)

comfort_df = df_to_dict(
    com_path,
    cols=["room_ids", "floor_level", "extreme_hot_satisf", "extreme_cold_satisf", "daylight_satisf", "sound_satisf", "air_quali", "comfort_satisfaction"],
    exclude_prefixes=exclude
)

delight_df = df_to_dict(
    del_path,
    cols=["room_ids", "floor_level", "hor_views_satisf", "sky_views_satisf", "green_views_satisf", "balcony_satisf", "space_size_satisf", "delight_satisfaction"],
    exclude_prefixes=exclude
)

social_df = df_to_dict(
    soc_path,
    cols=["room_ids", "floor_level", "social_amount_satisf", "social_distribution_satisf", "social_green_satisf", "social_satisfaction"],
    exclude_prefixes=exclude
)




thermal_reference = f"""
1. Thermal Comfort Standards

Thermal Comfort – Research Data Analysis (Numerical) showed that:

Apartments (Residents): Satisfaction most likely between 19.3 °C and 28.3 °C. Dissatisfaction most common below 16 °C or above 31.6 °C.

Offices (Workers): Satisfaction most likely between 20.8 °C and 26.2 °C. Dissatisfaction increases below 18.5 °C and above 28.3 °C.

Thermal Comfort – Research Data Analysis (Generalised)

Apartment residents reported satisfaction across a wider thermal range than office workers.

Residents are generally comfortable between cooler and warmer temperatures, while office workers indicated a narrower comfort band, preferring conditions closer to the mid-20s °C.

Both groups expressed dissatisfaction at lower extremes (below ~18 °C) and higher extremes (above ~30 °C).

2. Window Performance
Glazing Types:

Double Glazing: Two panes of glass with an air or gas-filled space between them, reducing heat transfer.

Triple Glazing: Adds a third pane, further enhancing insulation and energy efficiency.

U-Value:

The U-value measures the rate of heat transfer; lower values indicate better insulation.

Double Glazing: Typically has U-values ranging from 1.2 to 3.7 W/m²·K.

Triple Glazing: Often achieves U-values below 1.0 W/m²·K, offering superior thermal performance.

Low-Emissivity (Low-E) Coatings:

Low-E coatings are microscopically thin layers applied to glazing surfaces to reduce infrared radiation, minimizing heat transfer without compromising visible light transmission.

Solar Heat Gain Coefficient (SHGC):

SHGC measures how much solar radiation passes through a window.

Hot Climates: Lower SHGC (<0.35) reduces cooling loads.

Cold Climates: Higher SHGC (0.45–0.55) allows for passive solar heating.

3. Insulation Strategies
Wall Insulation:

Proper wall insulation is crucial for thermal comfort.

R-Values: For temperate climates, R-values above R-2.5-R-3.5 (m²·K/W) are considered minimal, with higher values recommended for colder regions.

Roof Insulation:

Roof assemblies significantly impact thermal buffering.

R-Values: Roof insulation should reach R-4.0 or higher in residential buildings, especially in top-floor dwellings.
"""


daylight_reference = f"""
Daylighting Design Guidelines for Enhanced Indoor Comfort and Energy Efficiency
Daylighting is the strategic use of natural light to illuminate building interiors. Its goal is to improve occupant well-being and reduce energy consumption by minimizing reliance on electric lighting. Effective daylighting design considers building orientation, glazing properties, window placement, interior reflectance, and glare control.

Standards and Guidelines

LEED v4 (Leadership in Energy and Environmental Design) awards credits for daylight access, visual comfort, and quality views.

Design Strategies
Window-to-Wall Ratio (WWR):

Maintain a balanced WWR to allow daylight penetration without excessive heat gain or glare.

A WWR between 20% and 40% is often recommended depending on orientation and climate zone.

Glazing Selection:

Use high-performance glazing with Visible Transmittance (VT) of 0.6–0.7 for good daylight entry.

Select low SHGC (Solar Heat Gain Coefficient) glass to minimize overheating in warmer climates.

Shading Devices:

Integrate fixed or dynamic shading systems such as overhangs, light shelves, and louvers.

These reduce direct glare and redirect light deeper into interior spaces.

Interior Reflectance:

Choose light-colored wall and ceiling finishes to improve light diffusion.

Keep partitions low or translucent to extend daylight access across open-plan layouts.
"""


acoustic_reference = f"""
Acoustic Design Guidelines for Building Performance
1. Acoustic Performance Standards
Effective acoustic design in buildings is guided by several key standards:

ANSI/ASA S12.60: Establishes acoustic performance criteria for classrooms, recommending maximum background noise levels of 35 dBA.

LEED and WELL Building Standards: Incorporate acoustic performance as a criterion for certification, promoting environments that support occupant comfort and wellbeing.

Recommended Noise Limits
International standards and guidelines define recommended indoor noise levels for comfort and health:

WHO Environmental Noise Guidelines (2018):

≤ 35 dBA in day-use rooms (e.g., offices, living rooms)

≤ 30 dBA in bedrooms at night

Outdoor façade exposure should not exceed 53 dBA Lden in urban environments (preferably <50 dBA)

ASHRAE Handbook – HVAC Applications (2023):

Recommends Noise Criteria (NC) levels of 25–35 for residential and office spaces

ISO 10052 / ISO 16283-3 (field measurement):

Establish procedures for evaluating outdoor-to-indoor sound level differences (DnT,w or D2m,nT,w)

Key Design Strategies 
1. Glazing Selection

Avoid single glazing in noise-exposed areas

Instead, use:

Double glazing (≥6/12/6 mm): STC 32–36
Laminated glass: significantly reduces higher-frequency noise
Asymmetric glazing (e.g., 6mm + 10mm): better at controlling low-frequency traffic noise

Laminated Glass: Incorporate laminated glazing, which includes a viscoelastic interlayer, to enhance sound insulation properties.

Double or Triple Glazing: Use multiple glazing layers with air or inert gas fills to improve acoustic performance.

Window-to-Wall Ratio: Optimize the size and placement of windows to balance natural light and acoustic insulation.

2. Wall Systems (Indicative R Values)
If walls are known to be lightweight (e.g., framed construction):
Add mass or specify external cladding with high Rw (≥45 dB)
Use staggered stud designs or double-skin façades (if known early)
Mineral Wool and Fiberglass Insulation: Effective for absorbing sound within wall cavities.

3. Openings and Penetrations
Limit use of fixed louvres or vents on exposed façades unless acoustically treated
Ensure window-to-wall ratio is minimized on noisy façades
"""


airquali_reference = f"""
Indoor relative humidity (RH) plays a critical role in shaping indoor air quality (IAQ) and associated health outcomes. A consistent body of literature indicates that maintaining RH between 40% and 60% is generally optimal for occupant comfort and health (Wolkoff, P., 2018). 
There is also evidence suggesting a link between moderate CO₂ elevations and Sick Building Syndrome symptoms (Tsai, D.H. et al. 2012). 
Maintaining levels below 1000 ppm is widely regarded as best practice to prevent fatigue, mitigate symptoms of Sick Building Syndrome, and support cognitive performance. 
Strategies such as demand-controlled ventilation and adequate outdoor air supply are essential to manage indoor CO₂ levels effectively.

Enable cross-ventilation through operable windows on opposing façades.
Size and locate openings to capture prevailing winds (especially in residential and mixed-mode spaces)

In spaces with non-operable windows, special attention must be given to ensuring adequate mechanical ventilation.
Provide space allowances for future mechanical ventilation, even in naturally ventilated schemes

3. Prevent RH Extremes through Envelope & Massing
Incorporate thermal mass and moderate glazing ratios to stabilise temperature and reduce condensation
Avoid unshaded large glazed surfaces that can drive up RH from condensation or limit fresh air access due to overheating concerns
Locate bathrooms/kitchens along external walls where moisture can be vented directly outdoors

References:

Wolkoff, P., 2018. Indoor air humidity, air quality, and health–An overview. International journal of hygiene and environmental health, 221(3), pp.376-390.

Tsai, D.H., Lin, J.S. and Chan, C.C., 2012. Office workers’ sick building syndrome and indoor carbon dioxide concentrations. Journal of occupational and environmental hygiene, 9(5), pp.345-351.
"""


delight_reference = f"""
Views of the outdoors:
Attractive outdoor views, particularly of the sky and natural scenes, support circadian wellbeing, visual comfort, cognitive performance, and overall satisfaction (Altomonte et al., 2020; Aries et al., 2015). 
Occupants generally prefer seating near windows, with higher satisfaction reported when window-to-wall ratios exceed 25% (Ko et al., 2022; Yeom et al., 2020). 
Larger ratios above ~65% bring limited additional benefit (Kim et al., 2021). 
Real-world experiences highlight that long-distance and dynamic views are important for daily wellbeing, which should be considered in early design through orientation, floorplate depth, and window placement.

Direct contact with nature:
Exposure to nature provides restorative effects, reduces stress, and improves wellbeing across diverse contexts (Capaldi et al., 2014; Howell and Passmore, 2013). 
In dense cities, private access can be enhanced through balconies, rooftop gardens, and vertical greenery systems. 
Balconies gained new significance during the pandemic, supporting residents’ wellbeing (Duarte et al., 2023; Nguyen et al., 2024). 
Their usability depends strongly on size, with spaces above ~10–15 m² used more frequently and perceived as more supportive (Smektała and Baborska-Narożny, 2022; Song et al., 2024).

Living-space size:
Adequate dwelling size is linked to family dynamics, privacy, social life, and educational performance (Foye, 2021). 
Overcrowding can negatively affect children’s learning and household relationships. In high-rise developments, minimised floor areas increase these risks, making early design decisions on layout, depth, and ceiling height critical. 
Embedding minimum space standards is essential for preventing overcrowding and supporting long-term wellbeing (Carmona et al., 2010).

References:

Altomonte, S., Allen, J., Bluyssen, P.M., Brager, G., Heschong, L., Loder, A., Schiavon, S., et al. (2020), Ten questions concerning well-being in the built environment, Building and Environment, Vol. 180, p. 106949, doi: 10.1016/j.buildenv.2020.106949.
Aries, M., Aarts, M. and van Hoof, J. (2015), Daylight and health: A review of the evidence and consequences for the built environment, Lighting Research & Technology, Vol. 47 No. 1, pp. 6–27, doi: 10.1177/1477153513509258.
Capaldi, C.A., Dopko, R.L. and Zelenski, J.M. (2014), The relationship between nature connectedness and happiness: a meta-analysis, Frontiers in Psychology, Vol. 5, doi: 10.3389/fpsyg.2014.00976.
Carmona, M., Gallent, N. and Sarkar, R. (2010), Space standards: the benefits. Housing Standards: evidence and research, No. Report prepared by University College London to CABE.
Duarte, C.C., Cortiços, N.D., Stefańska, A. and Stefańska, A. (2023), Home Balconies during the COVID-19 Pandemic: Future Architect’s Preferences in Lisbon and Warsaw, doi: https://doi.org/10.3390/app13010298.
Foye, C. (2021), “Social construction of house size expectations: testing the positional good theory and aspiration spiral theory using UK and German panel data, Housing Studies, Vol. 36 No. 9, pp. 1513–1532, doi: 10.1080/02673037.2020.1795086.
Nguyen, G. T. N., Tsaih, L. S.-J., Chen, J. C.-P., Tamariska, S. R., Coelho, A. M. F., & Kung, H.-Y. (2024), Balcony usage as a space to achieve human well-being during pandemic COVID-19, Journal of Asian Architecture and Building Engineering, pp. 1–13, doi: 10.1080/13467581.2024.2370408.
Howell, A.J. and Passmore, H.-A. (2013), The Nature of Happiness: Nature Affiliation and Mental Well-Being, in Keyes, C.L.M. (Ed.), Mental Well-Being, Springer Netherlands, Dordrecht, pp. 231–257, doi: 10.1007/978-94-007-5195-8_11.
Kim, S., Park, H. and Choo, S. (2021), Effects of Changes to Architectural Elements on Human Relaxation-Arousal Responses: Based on VR and EEG, International Journal of Environmental Research and Public Health, Vol. 18 No. 8, p. 4305, doi: 10.3390/ijerph18084305.
Ko, W.H., Kent, M.G., Schiavon, S., Levitt, B. and Betti, G. (2022), “A Window View Quality Assessment Framework, LEUKOS, Vol. 18 No. 3, pp. 268–293, doi: 10.1080/15502724.2021.1965889.
Smektała, M. and Baborska-Narożny, M. (2022), The use of apartment balconies: context, design and social norms, No. Buildings and Cities 3(1): 134-152, doi: https://doi.org/10.5334/bc.193.
Song, T., Xu, L., Zhao, F. and Du, Y. (2024), Healing properties of residential balcony: Assessment of the characteristics of balcony space in Shanghai’s collective housing, Journal of Building Engineering.
Yeom, S., Kim, H., Hong, T., Park, H.S. and Lee, D.-E. (2020), An integrated psychological score for occupants based on their perception and emotional response according to the windows’ outdoor view size, Building and Environment, Vol. 180, p. 107019, doi: 10.1016/j.buildenv.2020.107019.
"""


social_reference = f"""
Opportunities for social interactions:
Social connections are fundamental to wellbeing, with strong ties and supportive networks linked to better mental and physical health (Jordan, 2023; Umberson and Karas Montez, 2010). 
Loneliness and isolation are major health risks, and even minimal interactions or community-based initiatives can enhance wellbeing (Gunaydin et al., 2021; Leavell et al., 2019).
The built environment plays a central role: walkable neighbourhoods and accessible social spaces foster interaction (Weijs-Perrée et al., 2015), while poorly designed high-rise housing has been associated with isolation, stress, and reduced cohesion (Barros et al., 2019; Kearns et al., 2012). 
Recent studies highlight how older adults, students, and low-income residents are particularly vulnerable when social infrastructure is lacking (Nguyen et al., 2025, 2024).
Design strategies such as transitional or semi-private zones, vertical social pockets, and circulation areas can mitigate social overload while encouraging connection (Aw and Lim, 2016; Williams, 2005). 
The quality, size, and connectivity of communal areas are critical for fostering belonging, neighbourliness, and positive emotions (Barrie et al., 2023; Kleeman et al., 2023). 
Successful public and shared spaces work best when well-designed and actively supported by community programming, helping to strengthen cohesion and reduce loneliness in high-rise living.

References:

Barrie, H., McDougall, K., Miller, K. and Faulkner, D. (2023), The social value of public spaces in mixed-use high-rise buildings. Buildings and Cities, Vol. 4 No. 1, pp. 669–689, doi: 10.5334/bc.339.
Barros, P., Ng Fat, L., Garcia, L.M.T., Slovic, A.D., Thomopoulos, N., de Sá, T.H., Morais, P., et al. (2019), Social consequences and mental health outcomes of living in high-rise residential buildings and the influence of planning, urban design and architectural decisions: A systematic review”, Cities, Vol. 93, pp. 263–272, doi: 10.1016/j.cities.2019.05.015.
Aw, S. B. and Lim, P. I. (2016), The provision of vertical social pockets for better social interaction in high-rise living, PLANNING MALAYSIA, Special Issue IV, 163–180.
Gunaydin, G., Oztekin, H., Karabulut, D.H. and Salman-Engin, S. (2021), Minimal Social Interactions with Strangers Predict Greater Subjective Well-Being, Journal of Happiness Studies, Vol. 22 No. 4, pp. 1839–1853, doi: 10.1007/s10902-020-00298-6.
Jordan, M. (2023), The power of connection: Self-care strategies of social wellbeing, Journal of Interprofessional Education & Practice, Vol. 31, p. 100586, doi: 10.1016/j.xjep.2022.100586.
Kearns, A., Whitley, E., Mason, P. and Bond, L. (2012), Living the High Life? Residential, Social and Psychosocial Outcomes for High-Rise Occupants in a Deprived Context, Housing Studies, Vol. 27 No. 1, pp. 97–126, doi: 10.1080/02673037.2012.632080.
Kleeman, A., Giles-Corti, B., Gunn, L., Hooper, P. and Foster, S. (2023), The impact of the design and quality of communal areas in apartment buildings on residents neighbouring and loneliness, Cities, Vol. 133, p. 104126, doi: 10.1016/j.cities.2022.104126.
Leavell, M.A., Leiferman, J.A., Gascon, M., Braddick, F., Gonzalez, J.C. and Litt, J.S. (2019), Nature-Based Social Prescribing in Urban Settings to Improve Social Connectedness and Mental Well-being: a Review, Current Environmental Health Reports, Vol. 6 No. 4, pp. 297–308, doi: 10.1007/s40572-019-00251-7.
Nguyen, L., Van Den Berg, P., Kemperman, A. and Mohammadi, M. (2025), How does the layout of indoor communal spaces in low-income high-rise apartment buildings impact the social interactions between residents?, Cities & Health, pp. 1–24, doi: 10.1080/23748834.2025.2509739.
Nguyen, L.P., Van Den Berg, P.E., Kemperman, A.D. and Mohammadi, M. (2024), Social impacts of living in high-rise apartment buildings: The effects of buildings and neighborhoods, Journal of Urban Affairs, pp. 1–22, doi: 10.1080/07352166.2024.2311165.
Umberson, D. and Karas Montez, J. (2010), Social Relationships and Health: A Flashpoint for Health Policy, Journal of Health and Social Behavior, Vol. 51 No. 1_suppl, pp. S54–S66, doi: 10.1177/0022146510383501.
Weijs-Perrée, M., van den Berg, P., Arentze, T. and Kemperman, A. (2015), Factors influencing social satisfaction and loneliness: a path analysis, Journal of Transport Geography, Vol. 45, pp. 24–31, doi: 10.1016/j.jtrangeo.2015.04.004.
Williams, J. (2005), Designing Neighbourhoods for Social Interaction: The Case of Cohousing, Journal of Urban Design, Vol. 10 No. 2, pp. 195–227, doi: 10.1080/13574800500086998.
"""




avg, satisfied, neutral, dissatisfied, low_rooms, mean_comfort, mean_delight, mean_social = extract_wellbeing_summary(wb_path)

dimensions = {
    "Comfort": mean_comfort,
    "Delight": mean_delight,
    "Social": mean_social
}
sorted_dims = sorted(dimensions.items(), key=lambda x: x[1])
ranking_text = '\n'.join([f"{i+1}. {dim} ({score})" for i, (dim, score) in enumerate(sorted_dims)])
lowest_dimension = sorted_dims[0][0]
lowest_score = sorted_dims[0][1]

c_means, c_issues, c_factor_issue_counts, c_worst_factor = extract_comfort_summary(com_path)
dl_means, dl_issues, dl_factor_issue_counts, dl_worst_factor = extract_delight_summary(del_path)
social_means, social_issues, social_factor_issue_counts, social_worst_factor = extract_social_summary(soc_path)

delight_chunks = ["\n".join(dl_issues[:5])]

thermal_issues = [issue for issue in c_issues if "thermal" in issue.lower()]
daylight_issues = [issue for issue in c_issues if "daylight" in issue.lower()]
acoustic_issues = [issue for issue in c_issues if "acoustic" in issue.lower()]
airquali_issues = [issue for issue in c_issues if "air_quality" in issue.lower()]

materials = extract_materials_summary_enhanced(mat_path)




llm = ChatOpenAI(model="gpt-4", temperature=0.3, openai_api_key=key)



class ComfortAgent:
    def __init__(self, llm):
        self.prompt = PromptTemplate(
            input_variables=[
                "means_thermal", "means_daylight", "means_acoustic", "means_air_quality",
                "materials_glazing_type", "materials_glass_type", "materials_windows_u",
                "materials_shgc", "materials_window_noise_reduction", "materials_wall_r_insulation",
                "materials_wall_noise_reduction", "materials_roof_r_insulation", "materials_ground_r_insulation",
                "issues", "factor_issue_counts", "worst_factors", "chunks"
            ],
            template="""
You are an expert in indoor environmental quality and comfort design.

Scores for comfort factors (thermal, daylight, acoustic, air quality) range from 0 to 2:
- < 0.66 = dissatisfied
- 0.66 to <1.33 = neutral
- ≥ 1.33 = satisfied

---

## Materials Summary
Construction and glazing characteristics:
- Glazing: {materials_glazing_type}
- Glass type: {materials_glass_type}
- Window U-value: {materials_windows_u}
- SHGC: {materials_shgc}
- Window noise reduction: {materials_window_noise_reduction} dB
- Wall insulation R-value: {materials_wall_r_insulation}
- Wall noise reduction: {materials_wall_noise_reduction} dB
- Roof insulation R-value: {materials_roof_r_insulation}
- Ground insulation R-value: {materials_ground_r_insulation}

## Comfort Dimension
Mean Comfort Scores:
- Thermal: {means_thermal}
- Daylight: {means_daylight}
- Acoustic: {means_acoustic}
- Air Quality: {means_air_quality}

The lowest-performing factor(s): {worst_factors}, with these rooms rated dissatisfied: {issues}

---

Design feedback based on the evidence below:
{chunks}

Add placeholder: `comfort_images`
"""
        )
        self.chain = LLMChain(llm=llm, prompt=self.prompt)

    def run(self, data):
        return self.chain.run(**data)


class DelightAgent:
    def __init__(self, llm):
        self.prompt = PromptTemplate(
            input_variables=[
                "means_views", "means_balcony", "means_space_size",
                "issues", "factor_issue_counts", "worst_factors", "chunks"
            ],
            template="""
You are an expert in positive stimuli in architectural environments.

Delight scores range from 0 to 2:
- < 0.66 = dissatisfied
- 0.66 to <1.33 = neutral
- ≥ 1.33 = satisfied

## Delight Scores
Views: {means_views}
Balconies: {means_balcony}
Space Size: {means_space_size}

Lowest-performing factor(s): {worst_factors}

Detected issues:
{issues}

---

Based on the knowledge below, write a design feedback summary for Delight:
{chunks}

Add placeholder: `delight_images`
"""
        )
        self.chain = LLMChain(llm=llm, prompt=self.prompt)

    def run(self, data):
        return self.chain.run(**data)


class SocialAgent:
    def __init__(self, llm):
        self.prompt = PromptTemplate(
            input_variables=[
                "means_social_amount", "means_social_green", "means_social_distribution",
                "issues", "factor_issue_counts", "worst_factors", "chunks"
            ],
            template="""
You are an expert in social spatial design.

Social satisfaction scores range from 0 to 2:
- < 0.66 = dissatisfied
- 0.66 to <1.33 = neutral
- ≥ 1.33 = satisfied

## Social Scores
Amount: {means_social_amount}
Green Space: {means_social_green}
Distribution: {means_social_distribution}

Detected issues:
{issues}

Based on the references below, write a report and design suggestions:
{chunks}

Add placeholder: `social_images`
"""
        )
        self.chain = LLMChain(llm=llm, prompt=self.prompt)

    def run(self, data):
        return self.chain.run(**data)



# Coordinator Agent
class CoordinatorAgent:
    def __init__(self, agents):
        self.agents = agents
        self.results = {}

    def run(self, input_data):
        for name, agent in self.agents.items():
            if name in input_data:
                print(f"🔍 Running {name.title()}Agent...")
                self.results[name] = agent.run(input_data[name])
            else:
                print(f"⚠ No data provided for {name}")
        return self.results



comfort_agent = ComfortAgent(llm)
delight_agent = DelightAgent(llm)
social_agent = SocialAgent(llm)

coordinator = CoordinatorAgent({
    "comfort": comfort_agent,
    "delight": delight_agent,
    "social": social_agent
})

comfort_inputs = {
    "means_thermal": c_means["thermal"],
    "means_daylight": c_means["daylight"],
    "means_acoustic": c_means["acoustic"],
    "means_air_quality": c_means["air_quality"],
    "materials_glazing_type": materials["glazing_type"],
    "materials_glass_type": materials["glass_type"],
    "materials_windows_u": materials["windows_u"],
    "materials_shgc": materials["shgc"],
    "materials_window_noise_reduction": materials["window_noise_reduction"],
    "materials_wall_r_insulation": materials["wall_r_insulation"],
    "materials_wall_noise_reduction": materials["wall_noise_reduction"],
    "materials_roof_r_insulation": materials["roof_r_insulation"],
    "materials_ground_r_insulation": materials["ground_r_insulation"],
    "issues": "\n".join(c_issues),
    "factor_issue_counts": c_factor_issue_counts,
    "worst_factors": ", ".join(c_worst_factor),
    "chunks": thermal_reference + "\n\n" + daylight_reference + "\n\n" + acoustic_reference + "\n\n" + airquali_reference
}

delight_inputs = {
    "means_views": dl_means["views"],
    "means_balcony": dl_means["balcony"],
    "means_space_size": dl_means["space_size"],
    "issues": "\n".join(dl_issues),
    "factor_issue_counts": dl_factor_issue_counts,
    "worst_factors": ", ".join(dl_worst_factor),
    "chunks": delight_reference
}

social_inputs = {
    "means_social_amount": social_means["social_amount"],
    "means_social_green": social_means["social_green"],
    "means_social_distribution": social_means["social_distribution"],
    "issues": "\n".join(social_issues),
    "factor_issue_counts": social_factor_issue_counts,
    "worst_factors": ", ".join(social_worst_factor),
    "chunks": social_reference
}

results = coordinator.run({
    "comfort": comfort_inputs,
    "delight": delight_inputs,
    "social": social_inputs
})



collected_references = []


def add_report_section(doc, title, report_text, dimension, image_dir):
    global collected_references
    doc.add_page_break()
    doc.add_heading(title, level=1)

    lines = report_text.strip().split("\n")
    in_references_section = False

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.lower().startswith("references") or line.lower().startswith("## references"):
            in_references_section = True
            continue 

        if in_references_section:
            collected_references.append(line)
            continue  

        if f'{dimension}_images' in line:
            img1 = image_dir / f"{dimension}_factors.png"
            img2 = image_dir / f"{dimension}_satisfaction.png"
            table = doc.add_table(rows=1, cols=2)
            row = table.rows[0].cells
            if img1.exists():
                row[0].paragraphs[0].add_run().add_picture(str(img1), width=Inches(3))
            else:
                row[0].text = "[Missing image]"
            if img2.exists():
                row[1].paragraphs[0].add_run().add_picture(str(img2), width=Inches(3))
            else:
                row[1].text = "[Missing image]"
            doc.add_paragraph()
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            add_body_paragraph(doc, "• " + line[2:])
        else:
            add_body_paragraph(doc, line)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    pattern = re.compile(r"(\*\*[^\*]+\*\*|https?://\S+|\S+|\s+)")

    for part in pattern.findall(text):
        if not part.strip():
            p.add_run(part)  
        elif part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("http://") or part.startswith("https://"):
            run = p.add_run(part)
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.underline = True
            try:
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn

                r_id = doc.part.relate_to(
                    part,
                    reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                    is_external=True,
                )
                
                r = run._r
                hlink = OxmlElement("w:hyperlink")
                hlink.set(qn("r:id"), r_id)
                r.addprevious(hlink)
                hlink.append(r)
            except Exception as e:
                print(f"⚠ Could not embed link for {part}:", e)
        else:
            p.add_run(part)

    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)

def add_report_section(doc, title, report_text, dimension, image_dir):
    global collected_references
    doc.add_page_break()
    doc.add_heading(title, level=1)

    lines = report_text.strip().split("\n")
    in_references_section = False

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if f'{dimension}_images' in line:
            img1 = image_dir / f"{dimension}_factors.png"
            img2 = image_dir / f"{dimension}_satisfaction.png"
            table = doc.add_table(rows=1, cols=2)
            row = table.rows[0].cells
            if img1.exists():
                row[0].paragraphs[0].add_run().add_picture(str(img1), width=Inches(3))
            else:
                row[0].text = "[Missing image]"
            if img2.exists():
                row[1].paragraphs[0].add_run().add_picture(str(img2), width=Inches(3))
            else:
                row[1].text = "[Missing image]"
            doc.add_paragraph()
        elif line.startswith("## References"):

            in_references_section = True
        elif line.startswith("## "):
            in_references_section = False
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            add_body_paragraph(doc, "• " + line[2:])
        elif in_references_section:
            collected_references.append(line)
        else:
            add_body_paragraph(doc, line)




doc = Document()
doc.add_heading("Wellbeing by Design", 0)

add_report_section(doc, "Wellbeing", wellbeing_report, "wellbeing", wb_im_path)
add_report_section(doc, "Comfort", comfort_report, "comfort", com_im_path)
add_report_section(doc, "Delight", delight_report, "delight", del_im_path)
add_report_section(doc, "Social", social_report, "social", soc_im_path)
if collected_references:
    doc.add_page_break()
    doc.add_heading("References", level=1)
    for ref in sorted(set(collected_references)):
        add_body_paragraph(doc, ref)


output_path = out_f / "Wellbeing_Report.docx"
doc.save(output_path)
os.startfile(output_path)

